"""
Простой Markdown → DOCX конвертер для бизнес-решений.

Поддерживает: # h1 / ## h2 / ### h3 — заголовки;
              **bold** — жирный inline-текст;
              - / 1. — списки;
              | a | b | — Markdown-таблицы (одна шапка + строки);
              --- — разделитель;
              остальное — обычный параграф.

Не поддерживает: вложенные списки, ссылки, изображения. Для нашего
use-case (отчёт орестры) этого достаточно — стиль одинаковый с PDF.

Если python-docx не установлен (legacy-сервер до redeploy) — функция
возвращает False, caller использует fallback на markdown.
"""
from __future__ import annotations
import logging
import re

log = logging.getLogger(__name__)

_HEADING_RE = re.compile(r"^(#{1,4})\s+(.+?)\s*#*\s*$")
_BULLET_RE = re.compile(r"^\s*[-•*]\s+(.+)$")
_NUMBERED_RE = re.compile(r"^\s*\d+[\.\)]\s+(.+)$")
_TABLE_ROW_RE = re.compile(r"^\s*\|(.+)\|\s*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*[:\-]+\s*(\|\s*[:\-]+\s*)+\|?\s*$")
_HR_RE = re.compile(r"^\s*[-=*_]{3,}\s*$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_paragraph_with_bold(doc, text: str, *, style: str | None = None) -> None:
    """Добавить параграф с inline-bold (на месте **...**)."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        run = p.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def _flush_table(doc, header: list[str], rows: list[list[str]]) -> None:
    if not header or not rows:
        return
    cols = max(len(header), max((len(r) for r in rows), default=0))
    if cols < 1:
        return
    tbl = doc.add_table(rows=1 + len(rows), cols=cols)
    try:
        tbl.style = "Light Grid Accent 1"
    except Exception:
        pass
    # Шапка
    hdr_cells = tbl.rows[0].cells
    for i in range(cols):
        cell_text = header[i] if i < len(header) else ""
        hdr_cells[i].text = cell_text
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    # Строки
    for ri, row in enumerate(rows, start=1):
        for ci in range(cols):
            cell_text = row[ci] if ci < len(row) else ""
            tbl.rows[ri].cells[ci].text = cell_text


def markdown_to_docx(*, md_text: str, title: str, out_path: str,
                      subtitle: str = "") -> bool:
    """Конвертит markdown в DOCX. Возвращает True при успехе.

    `title` идёт первой строкой как Heading 1, потом subtitle (если есть)
    как параграф-курсив, потом разделитель, потом сам markdown.
    """
    try:
        import docx as _docx
        from docx.shared import Pt
    except Exception as e:
        log.warning(f"[docx] python-docx not available: {e}")
        return False

    try:
        doc = _docx.Document()
        # Заголовок
        h = doc.add_heading(title or "Бизнес-решение", level=0)
        if subtitle:
            sp = doc.add_paragraph(subtitle)
            sp.runs[0].italic = True
            sp.runs[0].font.size = Pt(11)
        doc.add_paragraph()  # отступ

        lines = (md_text or "").splitlines()
        i = 0
        n = len(lines)
        in_table = False
        tbl_header: list[str] = []
        tbl_rows: list[list[str]] = []

        def _close_table():
            nonlocal in_table, tbl_header, tbl_rows
            if in_table and tbl_header:
                _flush_table(doc, tbl_header, tbl_rows)
            in_table = False
            tbl_header = []
            tbl_rows = []

        while i < n:
            line = lines[i].rstrip()

            # Таблица: первая строка = шапка, вторая (если |---|---) = разделитель,
            # последующие — данные.
            if _TABLE_ROW_RE.match(line):
                if not in_table:
                    cols = [c.strip() for c in line.strip().strip("|").split("|")]
                    # Заглядываем на следующую строку — это разделитель?
                    if i + 1 < n and _TABLE_SEP_RE.match(lines[i + 1]):
                        in_table = True
                        tbl_header = cols
                        tbl_rows = []
                        i += 2  # пропускаем шапку и разделитель
                        continue
                    # Иначе — это просто параграф с | (бывает редко)
                else:
                    cols = [c.strip() for c in line.strip().strip("|").split("|")]
                    tbl_rows.append(cols)
                    i += 1
                    continue
            # Любая не-таблица строка — закрываем активную таблицу
            if in_table:
                _close_table()

            if not line:
                i += 1
                continue

            # Горизонтальная линия
            if _HR_RE.match(line):
                doc.add_paragraph("─" * 30)
                i += 1
                continue

            # Заголовки
            mh = _HEADING_RE.match(line)
            if mh:
                level = min(len(mh.group(1)), 4)
                doc.add_heading(mh.group(2).strip(), level=level)
                i += 1
                continue

            # Маркированный список
            mb = _BULLET_RE.match(line)
            if mb:
                _add_paragraph_with_bold(doc, mb.group(1).strip(),
                                          style="List Bullet")
                i += 1
                continue

            # Нумерованный список
            mn = _NUMBERED_RE.match(line)
            if mn:
                _add_paragraph_with_bold(doc, mn.group(1).strip(),
                                          style="List Number")
                i += 1
                continue

            # Обычный параграф
            _add_paragraph_with_bold(doc, line)
            i += 1

        _close_table()  # на случай если таблица была последней
        doc.save(out_path)
        return True
    except Exception as e:
        log.error(f"[docx] generation failed: {type(e).__name__}: {e}")
        return False
